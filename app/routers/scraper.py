from fastapi import APIRouter, Request, Depends, Form, HTTPException, status, BackgroundTasks
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.templating import Jinja2Templates
from sqlalchemy.orm import Session
import os
import re
import tempfile
import time
import uuid
import json
from typing import List, Dict, Any, Optional

import yt_dlp
from youtube_transcript_api import YouTubeTranscriptApi
from docx import Document
from docx.shared import Inches
from fpdf import FPDF
import google.generativeai as genai
import httpx

from app.database import get_db
from app.auth import get_current_user
from app.models import User
from app.config import settings

# Configurar la API Key de Gemini
if settings.GEMINI_API_KEY:
    genai.configure(api_key=settings.GEMINI_API_KEY)

router = APIRouter(prefix="/scraper", tags=["YouTube Scraper"])
templates = Jinja2Templates(directory="app/templates")

# Directorio temporal para almacenar descargas creadas
TEMP_DOWNLOADS_DIR = os.path.join(tempfile.gettempdir(), "herramientas_it_downloads")
os.makedirs(TEMP_DOWNLOADS_DIR, exist_ok=True)

def extract_video_id(url: str) -> str:
    """Extrae el Video ID de una URL de YouTube utilizando Regex"""
    pattern = r'(?:https?:\/\/)?(?:www\.)?(?:youtube\.com\/(?:[^\/\n\s]+\/\S+\/|(?:v|e(?:mbed)?)\/|\S*?[?&]v=)|youtu\.be\/)([a-zA-Z0-9_-]{11})'
    match = re.search(pattern, url)
    if not match:
        raise ValueError("URL de YouTube no válida.")
    return match.group(1)

def get_video_title(url: str, video_id: str) -> str:
    """Obtiene el título del video utilizando yt-dlp de forma rápida sin descargar el video"""
    ydl_opts = {
        'skip_download': True,
        'extract_flat': True,
        'quiet': True,
        'no_warnings': True,
    }
    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=False)
            return info.get('title', f"video_{video_id}")
    except Exception:
        return f"Transcripcion_YouTube_{video_id}"

def sanitize_filename(filename: str) -> str:
    """Remueve caracteres inválidos para nombres de archivo"""
    return re.sub(r'[\\/*?:"<>|]', "", filename).replace(" ", "_")

def clean_text_for_pdf(text: str) -> str:
    """Limpia caracteres no compatibles con fuentes latin-1 estándar de FPDF"""
    replacements = {
        '“': '"', '”': '"', '‘': "'", '’': "'", '—': '-', '–': '-',
        'á': 'a', 'é': 'e', 'í': 'i', 'ó': 'o', 'ú': 'u',
        'Á': 'A', 'É': 'E', 'Í': 'I', 'Ó': 'O', 'Ú': 'U',
        'ñ': 'n', 'Ñ': 'N', 'ü': 'u', 'Ü': 'U',
        '¿': '?', '¡': '!'
    }
    for orig, rep in replacements.items():
        text = text.replace(orig, rep)
    # Reemplazar tabulaciones por espacios para evitar fallos de cálculo de ancho
    text = text.replace('\t', '    ')
    # Codificar en latin-1 descartando caracteres no mapeables
    return text.encode('latin-1', 'replace').decode('latin-1')

def download_audio_via_ytdlp(video_id: str) -> str:
    """Descarga el audio de un video de YouTube como archivo m4a de forma optimizada sin requerir ffmpeg"""
    url = f"https://www.youtube.com/watch?v={video_id}"
    outtmpl = os.path.join(TEMP_DOWNLOADS_DIR, f"{video_id}.%(ext)s")
    
    ydl_opts = {
        'format': '140/m4a/bestaudio/best', # AAC m4a estándar de YouTube (no requiere transcoder/ffmpeg)
        'outtmpl': outtmpl,
        'quiet': True,
        'no_warnings': True,
    }
    
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        ydl.download([url])
        
    for ext in ['m4a', 'webm', 'mp3', 'aac', 'ogg', 'wav']:
        filepath = os.path.join(TEMP_DOWNLOADS_DIR, f"{video_id}.{ext}")
        if os.path.exists(filepath):
            return filepath
            
    raise Exception("No se pudo descargar el archivo de audio del video.")

def transcribe_and_format_audio(filepath: str, title: str, url: str) -> str:
    """Sube el archivo de audio a Gemini, realiza la transcripción y el formateo, y limpia el recurso remoto"""
    if not settings.GEMINI_API_KEY:
        raise Exception("GEMINI_API_KEY no está configurada en el servidor.")
        
    filename = os.path.basename(filepath)
    filesize = os.path.getsize(filepath)
    mime_type = "audio/mp4" # m4a es audio/mp4 o audio/x-m4a
    
    # 1. Iniciar la carga resumible
    init_url = f"https://generativelanguage.googleapis.com/upload/v1beta/files?key={settings.GEMINI_API_KEY}"
    headers = {
        "X-Goog-Upload-Protocol": "resumable",
        "X-Goog-Upload-Command": "start",
        "X-Goog-Upload-Header-Content-Length": str(filesize),
        "X-Goog-Upload-Header-Content-Type": mime_type,
        "Content-Type": "application/json",
    }
    json_data = {"file": {"display_name": filename}}
    
    with httpx.Client(timeout=600.0) as client:
        response = client.post(init_url, headers=headers, json=json_data)
        if response.status_code != 200:
            raise Exception(f"Fallo al iniciar carga en Gemini: {response.status_code} - {response.text}")
            
        upload_url = response.headers.get("X-Goog-Upload-URL")
        if not upload_url:
            raise Exception("No se recibió la URL de carga de Gemini.")
            
        # 2. Subir los bytes del archivo
        with open(filepath, "rb") as f:
            file_data = f.read()
            
        upload_headers = {
            "Content-Length": str(filesize),
            "X-Goog-Upload-Offset": "0",
            "X-Goog-Upload-Command": "upload, finalize",
        }
        
        response_upload = client.post(upload_url, headers=upload_headers, content=file_data)
        if response_upload.status_code != 200:
            raise Exception(f"Fallo al cargar bytes a Gemini: {response_upload.status_code} - {response_upload.text}")
            
        metadata = response_upload.json()
        file_resource_name = metadata.get("file", {}).get("name")
        if not file_resource_name:
            raise Exception("No se recibió el nombre del recurso de archivo de Gemini.")
            
        try:
            # 3. Esperar a que el archivo se procese (ACTIVE)
            file_info_url = f"https://generativelanguage.googleapis.com/v1beta/{file_resource_name}?key={settings.GEMINI_API_KEY}"
            for _ in range(60):
                info_resp = client.get(file_info_url)
                if info_resp.status_code == 200:
                    state = info_resp.json().get("state")
                    if state == "ACTIVE":
                        break
                    elif state == "FAILED":
                        raise Exception("El procesamiento del archivo falló en los servidores de Gemini.")
                time.sleep(2)
            else:
                raise Exception("Tiempo de espera agotado para el procesamiento del archivo en Gemini.")
                
            # 4. Generar el contenido estructurado
            gen_url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={settings.GEMINI_API_KEY}"
            prompt = (
                "Analiza de forma exhaustiva el siguiente audio de un video de soporte técnico, administración de sistemas o infraestructura.\n\n"
                "Tu tarea es actuar como un Ingeniero de Soporte Técnico Senior y generar un manual técnico o instructivo extremadamente detallado, completo y estructurado en español. "
                "Evita resúmenes cortos, explicaciones escuetas o simplificaciones. Queremos una guía exhaustiva donde se detalle cada paso con toda la información disponible en el audio.\n\n"
                "El documento final en Markdown debe estructurarse con las siguientes secciones:\n\n"
                "1. **Título Principal (#)**: Un título descriptivo y formal del procedimiento.\n"
                "2. **Introducción y Objetivos (##)**: Explica de manera detallada qué se logra con este procedimiento, la relevancia técnica y el contexto del sistema/herramienta descrita en el audio.\n"
                "3. **Requisitos Previos y Herramientas (##)**: Lista detallada de accesos necesarios, sistemas operativos compatibles, dependencias de software, herramientas físicas o digitales, comandos iniciales y configuraciones previas requeridas.\n"
                "4. **Guía Paso a Paso Detallada (##)**: Describe cronológicamente cada paso del procedimiento. Para cada paso:\n"
                "   - Explica el *qué*, el *cómo* y el *por qué* de la acción.\n"
                "   - Si el video menciona comandos de terminal, parámetros, scripts, configuraciones, switches de red, direcciones IP o sintaxis de código, "
                "escríbelos textualmente dentro de bloques de código en Markdown (```) indicando el lenguaje o shell (ej. bash, powershell, json, etc.). "
                "Añade comentarios o explicaciones de qué hace cada comando o parámetro.\n"
                "   - Detalla las respuestas del sistema esperadas, salidas de consola o confirmaciones visuales descritas.\n"
                "5. **Validación y Pruebas (##)**: Explica detalladamente cómo comprobar que el procedimiento se realizó con éxito (comandos de verificación, logs a revisar, pruebas de conectividad o funcionamiento).\n"
                "6. **Resolución de Problemas y Diagnósticos (##)**: Enumera los errores más comunes descritos o potenciales problemas que pueden surgir durante cada paso, y cómo solucionarlos.\n"
                "7. **Notas y Buenas Prácticas (##)**: Recomendaciones adicionales de seguridad, optimización o mantenimiento a largo plazo.\n\n"
                "Por favor, genera únicamente el documento Markdown estructurado y detallado, sin textos aclaratorios ni comentarios introductorios fuera del manual."
            )
            
            payload = {
                "contents": [
                    {
                        "parts": [
                            {
                                "file_data": {
                                    "mime_type": mime_type,
                                    "file_uri": f"https://generativelanguage.googleapis.com/v1beta/{file_resource_name}"
                                }
                            },
                            {
                                "text": prompt
                            }
                        ]
                    }
                ]
            }
            
            gen_resp = client.post(gen_url, json=payload, timeout=600.0)
            if gen_resp.status_code != 200:
                raise Exception(f"Fallo al generar contenido con Gemini: {gen_resp.status_code} - {gen_resp.text}")
                
            result = gen_resp.json()
            candidates = result.get("candidates", [])
            if not candidates:
                raise Exception("No se recibió respuesta del modelo Gemini.")
                
            parts = candidates[0].get("content", {}).get("parts", [])
            if not parts:
                raise Exception("No se recibió contenido de texto de Gemini.")
                
            return parts[0].get("text", "")
            
        finally:
            # 5. Eliminar el archivo de la API de Gemini Files
            try:
                del_url = f"https://generativelanguage.googleapis.com/v1beta/{file_resource_name}?key={settings.GEMINI_API_KEY}"
                client.delete(del_url)
            except Exception:
                pass

def generate_markdown_from_content(video_id: str, title: str, url: str, markdown_content: str) -> str:
    """Guarda la transcripción formateada en formato Markdown (.md)"""
    filename = sanitize_filename(f"{title}_{video_id}") + ".md"
    filepath = os.path.join(TEMP_DOWNLOADS_DIR, filename)
    
    with open(filepath, "w", encoding="utf-8") as f:
        if not markdown_content.strip().startswith("# "):
            f.write(f"# {title}\n\n")
            f.write(f"- **URL del Video**: [{url}]({url})\n")
            f.write(f"- **Video ID**: `{video_id}`\n\n")
        f.write(markdown_content)
        
    return filepath

def generate_docx_from_markdown(video_id: str, title: str, url: str, markdown_content: str) -> str:
    """Genera archivo Word (.docx) analizando el contenido estructurado del Markdown"""
    filename = sanitize_filename(f"{title}_{video_id}") + ".docx"
    filepath = os.path.join(TEMP_DOWNLOADS_DIR, filename)
    
    doc = Document()
    doc.add_heading(title, 0)
    
    p = doc.add_paragraph()
    p.add_run("URL del Video: ").bold = True
    p.add_run(url)
    p.add_run("\nID del Video: ").bold = True
    p.add_run(video_id)
    
    lines = markdown_content.splitlines()
    in_code_block = False
    code_content = []
    
    for line in lines:
        stripped = line.strip()
        
        if stripped.startswith("```"):
            if in_code_block:
                p_code = doc.add_paragraph()
                p_code.paragraph_format.left_indent = Inches(0.5)
                run = p_code.add_run("\n".join(code_content))
                run.font.name = 'Courier New'
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            continue
            
        if in_code_block:
            code_content.append(line)
            continue
            
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s+', stripped):
            content = re.sub(r'^\d+\.\s+', '', stripped)
            doc.add_paragraph(content, style='List Number')
        elif stripped == "":
            continue
        else:
            p_paragraph = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p_paragraph.add_run(part[2:-2])
                    run.bold = True
                else:
                    p_paragraph.add_run(part)
                    
    doc.save(filepath)
    return filepath

def generate_pdf_from_markdown(video_id: str, title: str, url: str, markdown_content: str) -> str:
    """Genera un reporte PDF con diseño limpio, moderno y profesional, imitando el estilo de Instructivo Modelo"""
    filename = sanitize_filename(f"{title}_{video_id}") + ".pdf"
    filepath = os.path.join(TEMP_DOWNLOADS_DIR, filename)
    
    class PDF(FPDF):
        def header(self):
            # Línea sutil de encabezado
            self.set_text_color(100, 116, 139) # slate-500
            self.set_font('Helvetica', '', 8)
            self.cell(self.epw / 2, 8, 'Instructivo de Asistencia Tecnica', border=False, align='L')
            self.cell(self.epw / 2, 8, 'Plataforma de Soporte IT', border=False, align='R')
            self.ln(6)
            self.set_draw_color(226, 232, 240) # slate-200
            self.set_line_width(0.2)
            self.line(self.l_margin, self.t_margin + 6, self.w - self.r_margin, self.t_margin + 6)
            self.ln(6)
            
        def footer(self):
            self.set_y(-15)
            self.set_draw_color(226, 232, 240)
            self.set_line_width(0.2)
            self.line(self.l_margin, self.h - 15, self.w - self.r_margin, self.h - 15)
            self.set_text_color(148, 163, 184) # slate-400
            self.set_font('Helvetica', 'I', 8)
            self.cell(0, 10, f'Pagina {self.page_no()} de {{nb}}', align='C')
            
    pdf = PDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    
    # 1. Título del Documento (Verde corporativo / Slate oscuro)
    pdf.set_font('Helvetica', 'B', 18)
    pdf.set_text_color(15, 23, 42) # slate-900
    pdf.multi_cell(pdf.epw, 9, clean_text_for_pdf(title))
    pdf.ln(3)
    
    # 2. Metadatos de origen
    pdf.set_font('Helvetica', '', 9)
    pdf.set_text_color(100, 116, 139) # slate-500
    
    pdf.write(5, clean_text_for_pdf("Canal de Origen: YouTube  |  "))
    pdf.write(5, clean_text_for_pdf("Video ID: "))
    pdf.set_font('Helvetica', 'B', 9)
    pdf.write(5, clean_text_for_pdf(video_id))
    pdf.ln(5)
    
    pdf.set_font('Helvetica', '', 9)
    pdf.write(5, clean_text_for_pdf("Enlace del Video: "))
    pdf.set_text_color(16, 185, 129) # emerald-500
    pdf.write(5, clean_text_for_pdf(url))
    pdf.ln(10)
    
    # Línea divisoria decorativa verde
    pdf.set_draw_color(16, 185, 129) # emerald-500
    pdf.set_line_width(1.0)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.ln(8)
    
    lines = markdown_content.splitlines()
    in_code_block = False
    
    for line in lines:
        stripped = line.strip()
        if not stripped and not in_code_block:
            pdf.ln(3)
            continue
            
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            pdf.set_font('Courier', '', 9)
            pdf.set_text_color(30, 41, 59) # slate-800
            pdf.set_fill_color(248, 250, 252) # slate-50 (fondo para bloque de código)
            pdf.multi_cell(pdf.epw, 5, clean_text_for_pdf(line), fill=True)
            continue
            
        # Títulos de nivel 1 (#)
        if stripped.startswith("# "):
            pdf.ln(5)
            pdf.set_font('Helvetica', 'B', 15)
            pdf.set_text_color(16, 185, 129) # emerald-500 (Verde para secciones)
            pdf.multi_cell(pdf.epw, 7, clean_text_for_pdf(stripped[2:]))
            pdf.ln(3)
            
        # Títulos de nivel 2 (##)
        elif stripped.startswith("## "):
            pdf.ln(4)
            pdf.set_font('Helvetica', 'B', 12)
            pdf.set_text_color(15, 23, 42) # slate-900 (Negro para subsecciones)
            pdf.multi_cell(pdf.epw, 6, clean_text_for_pdf(stripped[3:]))
            pdf.ln(2)
            
        # Títulos de nivel 3 (###)
        elif stripped.startswith("### "):
            pdf.ln(3)
            pdf.set_font('Helvetica', 'B', 11)
            pdf.set_text_color(71, 85, 105) # slate-600
            pdf.multi_cell(pdf.epw, 5, clean_text_for_pdf(stripped[4:]))
            pdf.ln(1.5)
            
        # Viñetas no numeradas (- o *)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            pdf.set_font('Helvetica', '', 10)
            pdf.set_text_color(51, 65, 85) # slate-700
            content = stripped[2:]
            # Reemplazar negritas Markdown **texto** en la viñeta
            content_clean = content.replace("**", "").replace("__", "")
            
            # Dibujar un punto de viñeta decente
            pdf.set_x(pdf.l_margin + 5)
            # Carácter de viñeta estándar (en latin-1, \x95 es el bullet point o usamos un guión limpio)
            bullet = chr(149) if hasattr(pdf, 'k') else "-"
            pdf.write(5, f"{bullet}  ")
            pdf.multi_cell(pdf.epw - 10, 5, clean_text_for_pdf(content_clean))
            
        # Lista numerada (ej. 1. o 2.)
        elif re.match(r'^\d+\.\s+', stripped):
            pdf.set_font('Helvetica', '', 10)
            pdf.set_text_color(51, 65, 85) # slate-700
            match = re.match(r'^(\d+\.\s+)(.*)', stripped)
            if match:
                num_part = match.group(1)
                text_part = match.group(2).replace("**", "").replace("__", "")
                pdf.set_x(pdf.l_margin + 2)
                pdf.write(5, clean_text_for_pdf(num_part))
                pdf.multi_cell(pdf.epw - 5, 5, clean_text_for_pdf(text_part))
            else:
                pdf.multi_cell(pdf.epw, 5, clean_text_for_pdf(stripped.replace("**", "")))
                
        # Párrafos normales
        else:
            pdf.set_font('Helvetica', '', 10)
            pdf.set_text_color(51, 65, 85) # slate-700
            clean_line = line.replace("**", "").replace("__", "")
            pdf.multi_cell(pdf.epw, 5, clean_text_for_pdf(clean_line))
            
    pdf.output(filepath)
    return filepath

@router.get("/", response_class=HTMLResponse)
async def get_scraper_page(request: Request, current_user: User = Depends(get_current_user)):
    """Renderiza la vista del Scraper de YouTube"""
    return templates.TemplateResponse(
        "scraper.html",
        {"request": request, "user": current_user}
    )

def save_task_status(task_id: str, status_data: Dict[str, Any]):
    """Guarda el estado de la tarea en un archivo JSON temporal para compartirlo entre workers de Gunicorn"""
    filepath = os.path.join(TEMP_DOWNLOADS_DIR, f"task_{task_id}.json")
    try:
        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(status_data, f, ensure_ascii=False)
    except Exception:
        pass

def get_task_status_data(task_id: str) -> Optional[Dict[str, Any]]:
    """Lee el estado de la tarea desde el archivo JSON temporal"""
    filepath = os.path.join(TEMP_DOWNLOADS_DIR, f"task_{task_id}.json")
    if os.path.exists(filepath):
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return None

def bg_process_transcription(task_id: str, video_url: str, video_id: str):
    """Tarea ejecutada en segundo plano para procesar el video sin bloquear el servidor web"""
    try:
        # 1. Obtener título del video
        save_task_status(task_id, {
            "status": "processing",
            "message": "Obteniendo información del video desde YouTube..."
        })
        title = get_video_title(video_url, video_id)
        
        # 2. Descargar flujo de audio
        save_task_status(task_id, {
            "status": "processing",
            "message": "Descargando flujo de audio optimizado (m4a)..."
        })
        audio_path = download_audio_via_ytdlp(video_id)
        
        # 3. Transcribir y formatear con Gemini
        save_task_status(task_id, {
            "status": "processing",
            "message": "Subiendo audio y transcribiendo con Gemini AI (esto puede tardar)..."
        })
        try:
            markdown_content = transcribe_and_format_audio(audio_path, title, video_url)
        finally:
            if os.path.exists(audio_path):
                os.remove(audio_path)
                
        if not markdown_content:
            raise Exception("No se pudo obtener la transcripción formateada de Gemini.")
            
        # 4. Generar archivos
        save_task_status(task_id, {
            "status": "processing",
            "message": "Generando documentos exportables (Markdown, Word, PDF)..."
        })
        md_path = generate_markdown_from_content(video_id, title, video_url, markdown_content)
        docx_path = generate_docx_from_markdown(video_id, title, video_url, markdown_content)
        pdf_path = generate_pdf_from_markdown(video_id, title, video_url, markdown_content)
        
        md_file = os.path.basename(md_path)
        docx_file = os.path.basename(docx_path)
        pdf_file = os.path.basename(pdf_path)
        
        save_task_status(task_id, {
            "status": "completed",
            "result": {
                "title": title,
                "video_id": video_id,
                "md_file": md_file,
                "docx_file": docx_file,
                "pdf_file": pdf_file
            }
        })
    except Exception as e:
        error_msg = str(e)
        if "Subtitles are disabled" in error_msg or "Could not find a transcript" in error_msg:
            friendly_error = "El video no tiene subtítulos o transcripciones disponibles para este ID."
        else:
            friendly_error = f"Error al procesar el video: {error_msg}"
            
        save_task_status(task_id, {
            "status": "failed",
            "error": friendly_error
        })

@router.post("/process", response_class=HTMLResponse)
async def process_youtube_url(
    request: Request,
    background_tasks: BackgroundTasks,
    video_url: str = Form(...),
    current_user: User = Depends(get_current_user)
):
    """Inicia el procesamiento del video en segundo plano y retorna de inmediato el spinner de polling"""
    video_url = video_url.strip()
    if not video_url:
        return templates.TemplateResponse(
            "partials/scraper_results.html",
            {"request": request, "error": "Por favor ingrese una URL válida."}
        )
        
    try:
        video_id = extract_video_id(video_url)
    except ValueError as e:
        return templates.TemplateResponse(
            "partials/scraper_results.html",
            {"request": request, "error": str(e)}
        )

    task_id = str(uuid.uuid4())
    save_task_status(task_id, {
        "status": "pending",
        "message": "Inicializando transcripción en segundo plano..."
    })
    
    # Iniciar la tarea en segundo plano sin bloquear el request
    background_tasks.add_task(bg_process_transcription, task_id, video_url, video_id)
    
    # Retornar inmediatamente un spinner que hace polling cada 3 segundos a la ruta de estado
    return HTMLResponse(
        content=f"""
        <div class="p-6 bg-white border border-slate-100 rounded-2xl custom-shadow flex flex-col items-center justify-center space-y-4 text-center"
             hx-get="/scraper/status/{task_id}" 
             hx-trigger="every 3s" 
             hx-target="#scraper-results-container" 
             hx-swap="innerHTML">
            <div class="relative w-12 h-12 flex items-center justify-center">
                <div class="absolute w-12 h-12 border-4 border-brand-100 rounded-full"></div>
                <div class="absolute w-12 h-12 border-4 border-t-brand-600 rounded-full animate-spin"></div>
            </div>
            <div>
                <span class="block text-sm font-bold text-slate-800">Iniciando proceso en segundo plano...</span>
                <span class="block text-xs text-slate-400 mt-1">Este proceso se ejecuta asíncronamente para evitar cortes por timeout. Puedes esperar aquí.</span>
            </div>
        </div>
        """
    )

@router.get("/status/{task_id}", response_class=HTMLResponse)
async def get_task_status(
    request: Request,
    task_id: str,
    current_user: User = Depends(get_current_user)
):
    """Endpoint de polling para comprobar el estado de una tarea y actualizar la UI mediante HTMX"""
    task = get_task_status_data(task_id)
    if not task:
        return HTMLResponse(
            content="<div class='p-4 bg-rose-50 border border-rose-100 rounded-xl text-rose-600 text-sm font-semibold'>Tarea no encontrada o expirada.</div>"
        )
        
    status_val = task.get("status")
    
    if status_val in ["pending", "processing"]:
        message = task.get("message", "Procesando...")
        return HTMLResponse(
            content=f"""
            <div class="p-6 bg-white border border-slate-100 rounded-2xl custom-shadow flex flex-col items-center justify-center space-y-4 text-center"
                 hx-get="/scraper/status/{task_id}" 
                 hx-trigger="every 3s" 
                 hx-target="#scraper-results-container" 
                 hx-swap="innerHTML">
                <div class="relative w-12 h-12 flex items-center justify-center">
                    <div class="absolute w-12 h-12 border-4 border-brand-100 rounded-full"></div>
                    <div class="absolute w-12 h-12 border-4 border-t-brand-600 rounded-full animate-spin"></div>
                </div>
                <div>
                    <span class="block text-sm font-bold text-slate-800">{message}</span>
                    <span class="block text-xs text-slate-400 mt-1">Procesando de forma segura sin límites de tiempo en el servidor...</span>
                </div>
            </div>
            """
        )
    elif status_val == "completed":
        res = task.get("result", {})
        return templates.TemplateResponse(
            "partials/scraper_results.html",
            {
                "request": request,
                "title": res.get("title"),
                "video_id": res.get("video_id"),
                "md_file": res.get("md_file"),
                "docx_file": res.get("docx_file"),
                "pdf_file": res.get("pdf_file"),
                "success": True
            }
        )
    else: # failed
        error_msg = task.get("error", "Error desconocido.")
        return templates.TemplateResponse(
            "partials/scraper_results.html",
            {"request": request, "error": error_msg}
        )

@router.get("/download/{filename}")
async def download_file(filename: str, current_user: User = Depends(get_current_user)):
    """Permite la descarga directa de un archivo de transcripción generado"""
    # Sanitizar el nombre del archivo para prevenir Path Traversal
    filename = os.path.basename(filename)
    filepath = os.path.join(TEMP_DOWNLOADS_DIR, filename)
    
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="El archivo solicitado ya no está disponible en el servidor.")
        
    return FileResponse(
        path=filepath,
        filename=filename,
        media_type="application/octet-stream"
    )
